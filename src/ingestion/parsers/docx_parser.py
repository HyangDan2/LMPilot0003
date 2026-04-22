from __future__ import annotations

from pathlib import Path

from src.ingestion.parsers.base import DocumentParser, ParserError
from src.models.schemas import ParsedDocument, Section


class DocxParser(DocumentParser):
    """MVP DOCX parser using python-docx paragraphs and heading styles."""

    file_type = "docx"

    def parse(self, path: Path) -> ParsedDocument:
        try:
            from docx import Document  # type: ignore[import-not-found]
        except Exception as exc:
            raise ParserError("DOCX parsing requires python-docx.") from exc

        try:
            document = Document(str(path))
        except Exception as exc:
            raise ParserError(f"Failed to open DOCX {path}: {exc}") from exc

        doc_id = self.doc_id(path)
        title = self.title_from_path(path)
        sections: list[Section] = []
        current_title = title
        current_level = 1
        current_lines: list[str] = []
        section_index = 0
        all_lines: list[str] = []

        def flush_section() -> None:
            nonlocal section_index, current_lines
            text = "\n".join(line for line in current_lines if line.strip()).strip()
            if not text and not current_title:
                return
            section_index += 1
            sections.append(
                Section(
                    section_id=f"{doc_id}-section-{section_index}",
                    title=current_title or f"Section {section_index}",
                    level=current_level,
                    text=text,
                    metadata={"parser": "docx"},
                )
            )
            current_lines = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            all_lines.append(text)
            style_name = paragraph.style.name if paragraph.style is not None else ""
            heading_level = _heading_level(style_name)
            if heading_level is not None:
                flush_section()
                current_title = text
                current_level = heading_level
            else:
                current_lines.append(text)

        flush_section()
        body = "\n".join(all_lines).strip()
        return ParsedDocument(
            doc_id=doc_id,
            file_path=str(path.resolve()),
            file_type=self.file_type,
            title=title,
            text=body,
            sections=sections,
            assets=[],
            metadata={"paragraph_count": len(document.paragraphs)},
        )


def _heading_level(style_name: str) -> int | None:
    normalized = style_name.strip().lower()
    if not normalized.startswith("heading"):
        return None
    parts = normalized.split()
    if len(parts) > 1 and parts[-1].isdigit():
        return max(1, int(parts[-1]))
    return 1

